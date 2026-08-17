from pathlib import Path
import sys

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path("/Users/dongyu/ClaudeCode/ESP32")
OUT_DIR = ROOT / "docs"
ASSET_DIR = ROOT / ".docx_qa" / "assets"
OUTPUT = OUT_DIR / "ESP32-S3_桌面智能控制台_功能规划_v1.0.docx"
FONT_PATH = Path("/System/Library/AssetsV2/com_apple_MobileAsset_Font8/86ba2c91f017a3749571a82f2c6d890ac7ffb2fb.asset/AssetData/PingFang.ttc")
TABLE_HELPER = Path("/Users/dongyu/.codex/plugins/cache/openai-primary-runtime/documents/26.813.12317/skills/documents/scripts")
sys.path.insert(0, str(TABLE_HELPER))
from table_geometry import apply_table_geometry


NAVY = "16324F"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "17202A"
MUTED = "65717D"
LIGHT = "F2F4F7"
BLUE_LIGHT = "E8EEF5"
CYAN_LIGHT = "EAF7FA"
GREEN = "2D7D62"
GREEN_LIGHT = "E9F5F0"
GOLD = "B7791F"
GOLD_LIGHT = "FFF7E6"
RED = "A23A3A"
RED_LIGHT = "FCECEC"
WHITE = "FFFFFF"
BLACK = "000000"

PAGE_W_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 100, "bottom": 100, "start": 120, "end": 120}


def rgb(hex_color):
    return RGBColor.from_string(hex_color)


def set_run_font(run, size=11, bold=None, color=INK, italic=None, latin="Arial Unicode MS", east_asia="Arial Unicode MS"):
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D7DCE2", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        el = tc_borders.find(tag)
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tc_borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:color"), color)


def set_cell_text(cell, text, *, bold=False, color=INK, size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(doc, headers, rows, widths, *, header_fill=BLUE_LIGHT, font_size=9.2, alignments=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.allow_autofit = False
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, color=NAVY, size=9.3,
                      align=(alignments[i] if alignments else WD_ALIGN_PARAGRAPH.LEFT))
        set_cell_shading(table.rows[0].cells[i], header_fill)
        set_cell_border(table.rows[0].cells[i])
    set_repeat_table_header(table.rows[0])
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            set_cell_text(cells[i], str(value), size=font_size,
                          align=(alignments[i] if alignments else WD_ALIGN_PARAGRAPH.LEFT))
            set_cell_border(cells[i])
    apply_table_geometry(table, widths, table_width_dxa=sum(widths),
                         indent_dxa=TABLE_INDENT_DXA, cell_margins_dxa=CELL_MARGINS)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(4)
    return table


def add_label_detail_table(doc, rows, *, label_width=1900, fill=LIGHT):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, detail in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], label, bold=True, color=NAVY, size=9.5)
        set_cell_text(cells[1], detail, color=INK, size=9.5)
        set_cell_shading(cells[0], fill)
        set_cell_border(cells[0])
        set_cell_border(cells[1])
    widths = [label_width, PAGE_W_DXA - label_width]
    apply_table_geometry(table, widths, table_width_dxa=PAGE_W_DXA,
                         indent_dxa=TABLE_INDENT_DXA, cell_margins_dxa=CELL_MARGINS)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_callout(doc, label, text, *, fill=BLUE_LIGHT, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, color=accent, size="10")
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.12
    r1 = p.add_run(f"{label}  ")
    set_run_font(r1, size=10.2, bold=True, color=accent)
    r2 = p.add_run(text)
    set_run_font(r2, size=10.2, color=INK)
    apply_table_geometry(table, [PAGE_W_DXA], table_width_dxa=PAGE_W_DXA,
                         indent_dxa=180, cell_margins_dxa={"top": 140, "bottom": 140, "start": 180, "end": 180})
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Arial Unicode MS")
    r_fonts.set(qn("w:hAnsi"), "Arial Unicode MS")
    r_fonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    r_pr.extend([r_fonts, color, underline])
    new_run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_page_field(paragraph):
    run = paragraph.add_run("第 ")
    set_run_font(run, size=8.5, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    r = OxmlElement("w:r")
    r.append(fld_char1)
    r.append(instr)
    r.append(fld_char2)
    paragraph._p.append(r)
    end = paragraph.add_run(" 页")
    set_run_font(end, size=8.5, color=MUTED)


def setup_numbering(doc):
    numbering = doc.part.numbering_part.element
    existing_abs = [int(el.get(qn("w:abstractNumId"))) for el in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    abs_id = max(existing_abs or [0]) + 1
    num_id = max(existing_num or [0]) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abs_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.extend([tabs, ind])
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Arial Unicode MS")
    r_fonts.set(qn("w:hAnsi"), "Arial Unicode MS")
    r_pr.append(r_fonts)
    lvl.extend([start, num_fmt, lvl_text, lvl_jc, p_pr, r_pr])
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_id = OxmlElement("w:abstractNumId")
    abstract_id.set(qn("w:val"), str(abs_id))
    num.append(abstract_id)
    numbering.append(num)
    return num_id


def set_bullet(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])


def add_bullet(doc, text, num_id, *, bold_prefix=None, color=INK):
    p = doc.add_paragraph()
    set_bullet(p, num_id)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.167
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=10.5, bold=True, color=NAVY)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, size=10.5, color=color)
    else:
        run = p.add_run(text)
        set_run_font(run, size=10.5, color=color)
    return p


def add_body(doc, text, *, bold_prefix=None, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=10.7, bold=True, color=NAVY)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, size=10.7, color=INK, italic=italic)
    else:
        r = p.add_run(text)
        set_run_font(r, size=10.7, color=INK, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    return p


def style_document(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)
    sec.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial Unicode MS"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Arial Unicode MS"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = sec.header
    hp = header.paragraphs[0]
    hp.text = ""
    hp.paragraph_format.space_after = Pt(0)
    hp.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    r1 = hp.add_run("ESP32-S3 桌面智能控制台")
    set_run_font(r1, size=8.5, bold=True, color=MUTED)
    hp.add_run("\t")
    r2 = hp.add_run("功能规划 v1.0")
    set_run_font(r2, size=8.5, color=MUTED)

    footer = sec.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    add_page_field(fp)


def add_title_page(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("产品功能规划")
    set_run_font(r, size=11, bold=True, color=GOLD)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("ESP32-S3 桌面智能控制台")
    set_run_font(r, size=27, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(22)
    r = p.add_run("4.3 英寸 · 800×480 · 触摸 · 低功耗蓝牙")
    set_run_font(r, size=13.5, color=MUTED)

    add_label_detail_table(doc, [
        ("目标设备", "微雪 ESP32-S3-Touch-LCD-4.3B，或同规格 4.3 英寸一体屏"),
        ("使用场景", "公司桌面信息中心、电脑控制台、个人工作流入口"),
        ("连接原则", "设备只走低功耗蓝牙；联网数据与系统控制由电脑端助手负责"),
        ("版本", "v1.0 · 2026-08-16"),
        ("状态", "功能范围规划与页面信息架构"),
    ], label_width=1750)

    add_callout(
        doc,
        "核心方向",
        "把屏幕做成一台常驻桌面的个人工作控制台：日常信息随时可见，常用操作一触即达，场景切换能够一次完成多项动作。",
        fill=CYAN_LIGHT,
        accent=DARK_BLUE,
    )

    add_heading(doc, "本版规划结论", 1)
    for text in [
        "采用九个功能页面，并保留全局状态栏、固定快捷入口与通知浮层。",
        "首版优先完成蓝牙连接、首页、系统监控、工作台、应用切换和亮灭屏逻辑。",
        "天气、市场、Codex、专注场景和自动化在首版稳定后逐步加入。",
        "高级玩法围绕情境配置、多动作、开发者状态、隐私控制与外接传感器展开。",
    ]:
        add_bullet(doc, text, doc._bullet_num_id)

    doc.add_page_break()


def load_font(size):
    return ImageFont.truetype(str(FONT_PATH), size=size)


def draw_centered(draw, box, text, font, fill):
    x0, y0, x1, y1 = box
    bb = draw.textbbox((0, 0), text, font=font)
    w, h = bb[2] - bb[0], bb[3] - bb[1]
    draw.text(((x0 + x1 - w) / 2, (y0 + y1 - h) / 2 - 2), text, font=font, fill=fill)


def create_page_map(path):
    img = Image.new("RGB", (1500, 780), "#F7F9FB")
    d = ImageDraw.Draw(img)
    f_title = load_font(36)
    f_box = load_font(26)
    f_small = load_font(21)
    d.text((60, 38), "页面体系与数据流", font=f_title, fill="#16324F")
    boxes = [
        (60, 130, 300, 240, "首页", "总览与提醒", "#EAF7FA"),
        (350, 130, 590, 240, "信息", "天气·市场·日程", "#FFF7E6"),
        (640, 130, 880, 240, "系统", "性能与网络", "#E9F5F0"),
        (930, 130, 1170, 240, "工作台", "快捷键与多动作", "#E8EEF5"),
        (1220, 130, 1460, 240, "应用", "启动与切换", "#FCECEC"),
        (205, 430, 445, 540, "Codex", "配额与令牌", "#E8EEF5"),
        (495, 430, 735, 540, "专注", "计时与会议", "#E9F5F0"),
        (785, 430, 1025, 540, "自动化", "场景与动作链", "#FFF7E6"),
        (1075, 430, 1315, 540, "设置", "设备与隐私", "#F2F4F7"),
    ]
    for x0, y0, x1, y1, title, subtitle, color in boxes:
        d.rounded_rectangle((x0, y0, x1, y1), radius=22, fill=color, outline="#C7D0D9", width=3)
        draw_centered(d, (x0, y0 + 12, x1, y0 + 66), title, f_box, "#16324F")
        draw_centered(d, (x0, y0 + 60, x1, y1 - 6), subtitle, f_small, "#65717D")
    for x in [300, 590, 880, 1170]:
        d.line((x + 8, 185, x + 42, 185), fill="#8DA2B5", width=4)
    for x in [325, 615, 905, 1195]:
        d.line((x, 255, x, 405), fill="#B8C4CE", width=3)
        d.line((x, 405, x + 0, 420), fill="#B8C4CE", width=3)
    d.rounded_rectangle((235, 650, 1265, 730), radius=20, fill="#16324F")
    draw_centered(d, (235, 650, 1265, 730), "全局层：蓝牙状态 · 数据新鲜度 · 通知 · 固定快捷入口 · 亮度与隐私", f_box, "#FFFFFF")
    img.save(path)


def draw_screen(draw, origin, title, cards, accent):
    ox, oy = origin
    x0, y0, x1, y1 = ox, oy, ox + 720, oy + 432
    draw.rounded_rectangle((x0, y0, x1, y1), radius=26, fill="#0E1621", outline="#35485C", width=4)
    draw.rectangle((x0 + 18, y0 + 16, x1 - 18, y0 + 56), fill="#172433")
    draw.text((x0 + 34, y0 + 22), title, font=load_font(22), fill="#FFFFFF")
    draw.text((x1 - 158, y0 + 22), "BLE ●  09:41", font=load_font(18), fill="#A9BAC8")
    positions = [(x0 + 24, y0 + 76, x0 + 348, y0 + 210), (x0 + 372, y0 + 76, x1 - 24, y0 + 210),
                 (x0 + 24, y0 + 228, x0 + 348, y0 + 362), (x0 + 372, y0 + 228, x1 - 24, y0 + 362)]
    for (bx0, by0, bx1, by1), (label, value) in zip(positions, cards):
        draw.rounded_rectangle((bx0, by0, bx1, by1), radius=18, fill="#172433", outline="#2C4154", width=2)
        draw.text((bx0 + 18, by0 + 16), label, font=load_font(18), fill="#94A8B8")
        draw.text((bx0 + 18, by0 + 56), value, font=load_font(28), fill=accent)
    draw.rounded_rectangle((x0 + 24, y1 - 54, x1 - 24, y1 - 16), radius=16, fill="#182A3B")
    draw.text((x0 + 50, y1 - 48), "首页     工作台     应用     Codex     更多", font=load_font(19), fill="#D7E1E8")


def create_screen_concepts(path):
    img = Image.new("RGB", (1600, 1040), "#F6F8FA")
    d = ImageDraw.Draw(img)
    d.text((70, 32), "代表性页面概念", font=load_font(38), fill="#16324F")
    draw_screen(d, (60, 110), "首页", [("时间", "09:41 周一"), ("天气", "28°C 多云"), ("系统", "CPU 34%"), ("Codex", "已用 31%")], "#69D5E7")
    draw_screen(d, (820, 110), "工作台", [("终端", "打开项目"), ("截屏", "区域截屏"), ("专注", "开始 50 分钟"), ("场景", "会议模式")], "#75D3A9")
    draw_screen(d, (60, 570), "Codex", [("主配额", "31%"), ("重置", "14:35"), ("今日令牌", "126K"), ("连续使用", "8 天")], "#9BB8FF")
    draw_screen(d, (820, 570), "专注与会议", [("专注计时", "42:18"), ("下个会议", "10:30"), ("麦克风", "已静音"), ("勿扰", "已开启")], "#F2C66D")
    img.save(path)


def add_figure(doc, image_path, caption, width=6.3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(str(image_path), width=Inches(width))
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(0)
    cp.paragraph_format.space_after = Pt(8)
    r = cp.add_run(caption)
    set_run_font(r, size=8.8, italic=True, color=MUTED)


def add_page_spec(doc, title, objective, components, interactions, states, priority, num_id):
    add_heading(doc, title, 2)
    add_body(doc, f"页面目标：{objective}", bold_prefix="页面目标：")
    add_label_detail_table(doc, [(label, detail) for label, detail in components], label_width=1750, fill=LIGHT)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("关键交互")
    set_run_font(r, size=10.3, bold=True, color=DARK_BLUE)
    for item in interactions:
        add_bullet(doc, item, num_id)
    add_body(doc, f"状态设计：{states}", bold_prefix="状态设计：")
    add_body(doc, f"实现优先级：{priority}", bold_prefix="实现优先级：")


def build_document():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    page_map = ASSET_DIR / "page_map.png"
    screen_concepts = ASSET_DIR / "screen_concepts.png"
    create_page_map(page_map)
    create_screen_concepts(screen_concepts)

    doc = Document()
    style_document(doc)
    doc._bullet_num_id = setup_numbering(doc)
    num_id = doc._bullet_num_id

    props = doc.core_properties
    props.title = "ESP32-S3 桌面智能控制台功能规划 v1.0"
    props.subject = "页面信息架构、功能范围、扩展玩法与开发优先级"
    props.author = "项目规划"
    props.keywords = "ESP32-S3, LVGL, BLE, 桌面控制台, Codex"

    add_title_page(doc)

    add_heading(doc, "一、产品定位与设计原则", 1)
    add_body(doc, "这台设备定位为公司桌面的个人信息与操作入口。ESP32 负责触摸、界面、状态展示和少量本地逻辑；电脑端助手负责联网、系统数据采集、应用控制、权限管理和 Codex 集成。")
    for text in [
        "抬眼可见：时间、天气、市场、系统与 Codex 状态在数秒内读懂。",
        "一触即达：高频操作控制在一次点击，复杂流程交给动作链。",
        "情境驱动：前台应用、时间、会议状态和用户选择都可以切换页面配置。",
        "公司可用：蓝牙绑定、动作白名单、敏感信息默认隐藏、设备端不保存账号密钥。",
        "渐进扩展：关键功能在无存储卡时可用，主题、图标、日志和截图可放入 TF 卡。",
    ]:
        add_bullet(doc, text, num_id)
    add_figure(doc, page_map, "图 1：九个功能页面通过全局状态层连接", width=6.35)
    add_callout(doc, "交互基线", "屏幕连接成功后渐亮；心跳丢失后进入短暂重连等待；持续断开时渐暗并关闭背光。", fill=GREEN_LIGHT, accent=GREEN)

    add_heading(doc, "二、全局导航与交互", 1)
    add_table(doc,
              ["区域", "建议高度", "承担内容", "交互"],
              [
                  ("顶部状态栏", "38 px", "时间、蓝牙、数据新鲜度、隐私状态", "点击进入对应状态详情"),
                  ("主内容区", "390 px", "当前页面卡片、图表或按钮", "点击、长按、滑动"),
                  ("底部导航", "52 px", "首页、工作台、应用、Codex、更多", "固定入口与页面指示"),
              ],
              [1700, 1300, 3860, 2500],
              alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT])
    add_heading(doc, "建议手势", 2)
    for text in [
        "单击：执行动作或进入详情；执行后显示成功、失败或进行中反馈。",
        "长按：打开动作说明、二次确认或编辑入口。",
        "双击：留给用户自定义的高频动作，首版默认关闭。",
        "左右滑动：在同一信息组内翻页，例如市场指数和系统曲线。",
        "从顶部下滑：打开亮度、勿扰、隐私与连接状态面板。",
    ]:
        add_bullet(doc, text, num_id)
    add_figure(doc, screen_concepts, "图 2：首页、工作台、Codex、专注页面的视觉方向", width=6.35)

    doc.add_page_break()
    add_heading(doc, "三、页面功能规划", 1)
    add_page_spec(doc, "页面一：首页仪表盘", "用最少视线停留完成工作状态总览。",
                  [
                      ("核心区", "大号时间、日期、星期；工作时段可显示下一项日程。"),
                      ("信息卡", "天气、温度、空气质量、今日高低温。"),
                      ("市场卡", "两到四个自选指数，显示涨跌幅与交易状态。"),
                      ("状态卡", "CPU、内存、网络与 Codex 配额摘要。"),
                      ("提示区", "会议倒计时、专注计时、异常提醒和数据过期提示。"),
                  ],
                  ["点击任意卡片进入对应详情页。", "长按时间区域切换极简时钟和完整仪表盘。", "双击首页空白区域执行用户指定的全局动作。"],
                  "无网络数据时保留最后一次结果并显示更新时间；电脑助手未完成认证时只显示本地时间与连接提示。",
                  "P0；第一批可用页面。", num_id)

    add_page_spec(doc, "页面二：信息中心", "集中呈现天气、市场和日程，减少首页拥挤。",
                  [
                      ("天气分区", "当前天气、小时趋势、未来三天、降雨概率、空气质量。"),
                      ("市场分区", "自选指数、汇率或数字资产；开盘时提高刷新频率。"),
                      ("日程分区", "下一场会议、今日剩余安排和会议链接状态。"),
                      ("切换方式", "顶部标签或横向滑动切换天气、市场、日程。"),
                  ],
                  ["点击指数进入当日走势图。", "长按项目进入自选列表。", "点击会议卡可在电脑上打开会议应用或日历详情。"],
                  "数据提供方失败时显示缓存与来源状态；价格信息明确标记延迟属性。",
                  "P1；首页稳定后加入。", num_id)

    doc.add_page_break()
    add_page_spec(doc, "页面三：系统监控", "持续观察电脑运行状态，并在异常时给出易懂提示。",
                  [
                      ("性能", "CPU、内存、磁盘占用和运行时间。"),
                      ("网络", "上行、下行、延迟和连接状态。"),
                      ("电源", "电池电量、充电状态和预计续航。"),
                      ("曲线", "最近 60 秒或 5 分钟趋势。"),
                      ("扩展", "温度、风扇和进程排行依据系统权限逐步开放。"),
                  ],
                  ["点击指标放大曲线。", "长按指标设置告警阈值。", "异常时跳转到电脑端助手查看原因。"],
                  "采样暂停、权限缺失和传感器不可用需显示独立状态，避免用零值掩盖异常。",
                  "P0；与蓝牙数据链路同步开发。", num_id)

    add_page_spec(doc, "页面四：工作台", "承载固定快捷键、宏动作和用户自定义工作流。",
                  [
                      ("布局", "默认 4×2 大按钮；可切换 3×2 大触控模式。"),
                      ("首批动作", "终端、截屏、静音、专注、锁屏、打开 Codex、打开项目、会议模式。"),
                      ("多动作", "一个按钮按顺序执行多个动作，可插入等待和条件判断。"),
                      ("固定入口", "首页、隐私键和返回键可以跨页面固定。"),
                  ],
                  ["点击立即执行低风险动作。", "关机、退出应用等高影响动作要求长按确认。", "长按空白按钮进入电脑端编辑器。"],
                  "每个动作均显示进行中、成功和失败状态；断连时按钮进入不可用态。",
                  "P0；首版先实现八个固定动作，编辑器进入 P1。", num_id)

    doc.add_page_break()
    add_page_spec(doc, "页面五：应用切换器", "替代频繁寻找 Dock 或使用键盘组合键。",
                  [
                      ("应用网格", "常用、最近、运行中三种排序方式。"),
                      ("状态", "突出当前前台应用，并标记有通知或正在播放的应用。"),
                      ("配置", "电脑端同步应用名称、图标与允许执行的动作。"),
                      ("情境页面", "前台应用变化时，可自动切换到该应用对应的快捷配置。"),
                  ],
                  ["单击激活应用。", "再次点击当前应用可在其窗口间切换。", "长按打开退出、隐藏、固定到首页等操作。"],
                  "应用图标缺失时使用通用图标；退出应用等操作必须有确认或白名单。",
                  "P0；先支持收藏应用，运行中列表进入 P1。", num_id)

    add_page_spec(doc, "页面六：Codex 面板", "展示账户用量与开发节奏，并为后续任务状态集成保留位置。",
                  [
                      ("配额", "当前配额桶使用百分比、窗口长度和重置时间。"),
                      ("令牌", "累计令牌、每日令牌、峰值和连续使用天数。"),
                      ("提醒", "接近阈值时改变颜色，并显示预计重置时间。"),
                      ("任务状态", "当电脑助手通过同一 App Server 管理任务时，可扩展运行、等待、完成状态。"),
                  ],
                  ["点击配额查看各配额桶。", "滑动查看最近七天令牌趋势。", "点击 Codex 快捷键切换到桌面应用。"],
                  "账户认证不可用时显示引导；配额窗口由接口动态返回，不写死固定周期。OpenAI Docs 已确认配额读取、更新通知和令牌统计接口可用。[3]",
                  "P1；账户用量优先，任务状态列为 P2。", num_id)

    doc.add_page_break()
    add_page_spec(doc, "页面七：专注与会议", "把计时、勿扰、麦克风和会议状态放在同一处。",
                  [
                      ("专注", "25/50/90 分钟计时、休息提示、今日累计专注时长。"),
                      ("会议", "下一场会议、倒计时、会议时长、麦克风状态。"),
                      ("快捷控制", "静音、勿扰、打开会议、记一条速记、延长专注。"),
                      ("极简显示", "专注开始后可隐藏其他信息，只保留计时和一个退出按钮。"),
                  ],
                  ["点击计时器开始或暂停。", "长按结束专注，降低误触。", "会议开始时可自动切换到会议页面。"],
                  "日历、麦克风和会议应用权限分别管理；没有权限时仍可使用纯计时模式。",
                  "P1；先做计时和勿扰，再接日历与会议应用。", num_id)

    add_page_spec(doc, "页面八：自动化与场景", "用一个场景完成一组动作，并让用户清楚看到执行进度。",
                  [
                      ("到岗场景", "连接设备、打开工作应用、同步数据、恢复工作音量。"),
                      ("专注场景", "打开项目与终端、开启勿扰、启动计时、切换专注页面。"),
                      ("会议场景", "打开会议应用、静音系统通知、显示议程与计时。"),
                      ("演示场景", "关闭敏感通知、调整亮度、打开演示文档。"),
                      ("离岗场景", "保存状态、关闭屏幕；锁屏动作由用户单独启用。"),
                  ],
                  ["点击场景查看动作清单，再次确认后执行。", "执行过程中逐项显示状态。", "失败时允许跳过、重试或停止后续动作。"],
                  "自动场景默认关闭；到岗与离岗动作需要独立授权，避免电脑短暂断连触发意外操作。",
                  "P1；基础动作链先行，条件判断与自动触发进入 P2。", num_id)

    add_page_spec(doc, "页面九：设置与诊断", "集中管理连接、显示、隐私、数据和维护。",
                  [
                      ("显示", "亮度、主题、熄屏等待、夜间调暗和极简时钟。"),
                      ("连接", "设备名称、绑定电脑、重新配对、信号和心跳。"),
                      ("隐私", "窗口名、通知、剪贴板、日历、截图与日志开关。"),
                      ("数据", "天气城市、自选指数、刷新频率和缓存状态。"),
                      ("维护", "固件版本、内存、TF 卡、错误日志、重启和恢复设置。"),
                  ],
                  ["危险操作进入二级页面并长按确认。", "权限问题可一键跳转电脑端助手。", "诊断页支持导出脱敏日志。"],
                  "关键设置保存在 NVS；大型资源和历史日志可放 TF 卡；无卡时自动降级。",
                  "P0 提供基础诊断，完整设置页进入 P1。", num_id)

    doc.add_page_break()
    add_heading(doc, "四、扩展玩法库", 1)
    add_callout(doc, "灵感来源", "成熟控制台产品常用页面、配置、文件夹、多动作、智能情境和固定动作来组织复杂能力；本项目可以吸收这些交互思路，同时保持自己的信息仪表盘特色。[4][5]", fill=GOLD_LIGHT, accent=GOLD)
    idea_rows = [
        ("一键开工", "打开项目、终端、浏览器与 Codex，开启勿扰并启动 50 分钟计时。", "P1"),
        ("一键会议", "打开会议应用、隐藏敏感通知、显示议程、启用会议计时。", "P1"),
        ("智能配置", "根据当前前台应用自动切换到设计、开发、写作或会议按钮布局。", "P2"),
        ("构建哨兵", "显示 Git 分支、未提交修改、测试结果、持续集成状态和本地服务端口。", "P2"),
        ("Codex 状态灯塔", "显示配额、任务运行、等待确认和完成状态；任务态依赖同一 App Server 集成。", "P2"),
        ("剪贴板货架", "展示用户主动固定的文本片段，点击后粘贴到当前应用。默认关闭自动采集。", "P2"),
        ("隐私场景键", "隐藏敏感窗口、关闭通知预览、静音麦克风，或直接锁定电脑。", "P1"),
        ("晨间简报", "到岗后显示天气、日程、市场、待办与 Codex 配额摘要。", "P2"),
        ("下班回顾", "展示专注时长、完成任务、令牌使用与明日第一项安排。", "P2"),
        ("像素伙伴", "用角色表情反映连接、专注、系统压力和 Codex 配额，作为可选主题。", "P3"),
        ("休息小游戏", "加入 2048、贪吃蛇、记忆翻牌或午餐选择器，限制在休息模式。", "P3"),
        ("二维码卡片", "显示个人名片、临时网页或会议链接二维码，避免存放公司无线网络密码。", "P3"),
        ("桌面环境站", "通过 I2C 外接温湿度、二氧化碳或环境光传感器。", "P3"),
        ("实体控制扩展", "外接旋钮、机械按键、脚踏开关或状态灯，提升盲操作体验。", "P3"),
        ("家庭与工作室控制", "由电脑端助手连接 Home Assistant，控制灯光、插座或场景。", "P3"),
    ]
    add_table(doc, ["玩法", "体验", "建议阶段"], idea_rows, [1800, 6260, 1300], header_fill=GOLD_LIGHT,
              font_size=8.8, alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER])

    add_heading(doc, "最有辨识度的三项玩法", 2)
    for text in [
        "场景控制：到岗、专注、会议、演示和离岗形成一套每天会使用的工作节奏。",
        "开发者仪表盘：系统、Git、构建、服务和 Codex 状态共同组成个人开发控制台。",
        "像素伙伴主题：把连接、专注和系统压力转成有生命感的视觉反馈，增加长期摆放意愿。",
    ]:
        add_bullet(doc, text, num_id)

    doc.add_page_break()
    add_heading(doc, "五、功能优先级", 1)
    priority_rows = [
        ("P0 · 可用闭环", "显示触摸、蓝牙绑定、亮灭屏、首页、系统、工作台、收藏应用、基础设置", "能够每天使用"),
        ("P1 · 完整产品", "信息中心、Codex 用量、专注、会议、动作链、隐私键、完整诊断", "形成稳定个人工作台"),
        ("P2 · 深度联动", "智能配置、运行中应用、Git/构建、Codex 任务态、剪贴板货架、自动触发", "融入开发工作流"),
        ("P3 · 个性扩展", "像素伙伴、小游戏、二维码、环境传感器、旋钮与家庭自动化", "形成个人特色"),
    ]
    add_table(doc, ["阶段", "范围", "阶段目标"], priority_rows, [2100, 5060, 2200], header_fill=BLUE_LIGHT,
              font_size=9.1)

    add_heading(doc, "建议首版固定按钮", 2)
    add_table(doc, ["位置", "动作", "短按", "长按"], [
        ("1", "应用切换", "打开收藏应用页", "打开运行中应用页"),
        ("2", "终端", "激活终端", "打开当前项目终端"),
        ("3", "截屏", "区域截屏", "整屏或窗口截屏"),
        ("4", "Codex", "激活 Codex", "打开 Codex 面板"),
        ("5", "静音", "切换静音", "打开音量面板"),
        ("6", "专注", "开始默认计时", "选择时长"),
        ("7", "会议", "执行会议场景", "预览动作清单"),
        ("8", "锁屏", "显示确认", "长按立即锁屏"),
    ], [850, 1800, 3250, 3460], header_fill=GREEN_LIGHT, font_size=8.9)

    doc.add_page_break()
    add_heading(doc, "六、数据、通信与安全边界", 1)
    add_heading(doc, "数据刷新建议", 2)
    add_table(doc, ["数据", "刷新方式", "默认频率", "断线策略"], [
        ("蓝牙心跳", "电脑助手主动发送", "3 秒", "15 秒无心跳进入熄屏流程"),
        ("系统状态", "电脑本地采样", "1～2 秒", "保留最后值并标记暂停"),
        ("天气", "电脑端联网", "10～30 分钟", "保留缓存与更新时间"),
        ("市场", "电脑端联网", "开盘 15～60 秒", "休市降低频率"),
        ("日历", "电脑本地权限", "5 分钟或事件驱动", "隐藏敏感字段"),
        ("Codex", "本地 App Server", "事件驱动，60 秒兜底", "认证失败显示状态"),
    ], [1800, 2400, 2200, 2960], header_fill=LIGHT, font_size=8.8)

    add_heading(doc, "蓝牙协议建议", 2)
    for text in [
        "ESP32 作为低功耗蓝牙外设，电脑助手负责扫描、连接、认证和重连。ESP32-S3 支持低功耗蓝牙，经典蓝牙不可用。[2]",
        "自定义 GATT 服务分为状态通知、数据更新、动作请求、动作回执和大文件分片。",
        "首版使用带版本号、序号和消息类型的 JSON；图标和截图采用独立二进制分片。",
        "每个触摸动作都有回执，屏幕必须能区分进行中、成功、失败和超时。",
        "常用媒体键可在后期增加蓝牙 HID，系统级复杂动作继续由电脑助手完成。",
    ]:
        add_bullet(doc, text, num_id)

    add_heading(doc, "安全与公司使用", 2)
    for text in [
        "启用配对、绑定和白名单，只允许已授权电脑连接。",
        "设备端不保存天气、市场、OpenAI 或公司系统的账号密钥。",
        "电脑助手维护动作白名单，ESP32 只发送动作编号与有限参数。",
        "剪贴板、窗口标题、通知正文、日历详情和截图功能默认关闭，逐项授权。",
        "日志进行脱敏，避免记录窗口正文、文件内容、令牌和个人信息。",
        "锁屏、退出应用、运行脚本等高影响动作使用长按确认或电脑端二次确认。",
    ]:
        add_bullet(doc, text, num_id)

    add_heading(doc, "存储策略", 2)
    add_label_detail_table(doc, [
        ("板载闪存", "固件、开机页、配对页、关键字体子集、默认主题和恢复资源。"),
        ("PSRAM", "LVGL 帧缓冲、图片缓存、图表数据和临时传输缓冲。"),
        ("TF 卡", "完整中文字体、应用图标、主题、截图、日志和可选动画。"),
        ("降级原则", "未插卡或读取失败时仍能进入首页、工作台、连接和设置。"),
    ])
    add_body(doc, "微雪官方资料确认目标板具备 16MB 闪存、8MB PSRAM 和 TF 卡槽。[1]")

    doc.add_page_break()
    add_heading(doc, "七、实施路线与验收标准", 1)
    roadmap_rows = [
        ("阶段 A", "硬件底座", "屏幕、触摸、背光、RTC、PSRAM、TF 卡和基础 LVGL"),
        ("阶段 B", "蓝牙闭环", "绑定、重连、心跳、亮灭屏、数据与动作回执"),
        ("阶段 C", "核心页面", "首页、系统、工作台、收藏应用和设置"),
        ("阶段 D", "电脑助手", "系统采样、应用控制、权限、动作白名单和配置"),
        ("阶段 E", "在线信息", "天气、市场、日历与缓存策略"),
        ("阶段 F", "Codex 与场景", "用量、令牌、专注、会议和动作链"),
        ("阶段 G", "产品化", "安全、日志、异常恢复、主题、性能与长期运行"),
    ]
    add_table(doc, ["阶段", "主题", "交付内容"], roadmap_rows, [1450, 1850, 6060], header_fill=BLUE_LIGHT, font_size=9.0)

    add_heading(doc, "首版验收标准", 2)
    for text in [
        "设备启动后可以完成配对；授权电脑靠近并运行助手时自动连接。",
        "完成认证和心跳后屏幕渐亮；连接持续丢失后屏幕自动关闭。",
        "首页系统数据连续更新；断线、权限缺失和数据过期均有明确提示。",
        "八个固定动作都有回执，失败不会让界面卡住。",
        "应用切换、终端、截屏、静音、专注、Codex 和锁屏均可稳定执行。",
        "连续运行一个工作日无明显内存增长、触摸失效或频繁重连。",
        "未插 TF 卡时仍能正常启动和使用 P0 功能。",
        "敏感数据不会出现在设备日志和默认界面。",
    ]:
        add_bullet(doc, text, num_id)

    add_callout(doc, "下一步", "进入阶段 A 前，确认开发板准确型号、公司电脑系统和首批需要控制的应用清单。随后建立固件、电脑助手、协议、资源与文档目录。", fill=CYAN_LIGHT, accent=DARK_BLUE)

    add_heading(doc, "八、参考资料", 1)
    sources = [
        ("[1] 微雪 ESP32-S3-Touch-LCD-4.3B 官方文档", "https://docs.waveshare.com/ESP32-S3-Touch-LCD-4.3B"),
        ("[2] 乐鑫 ESP32-S3 Bluetooth LE 官方说明", "https://docs.espressif.com/projects/esp-idf/en/stable/esp32s3/api-guides/ble/overview.html"),
        ("[3] OpenAI Docs：Codex App Server", "https://learn.chatgpt.com/docs/app-server"),
        ("[4] Elgato Stream Deck Profiles 官方文档", "https://docs.elgato.com/stream-deck/profiles/getting-started/"),
        ("[5] Elgato Stream Deck Multi Actions 官方说明", "https://help.elgato.com/hc/en-us/articles/360027960912-Elgato-Stream-Deck-Multi-Actions"),
        ("[6] Apple 支持：在 Mac 命令行运行快捷指令", "https://support.apple.com/en-ca/guide/shortcuts-mac/apd455c82f02/mac"),
    ]
    for label, url in sources:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(7)
        add_hyperlink(p, label, url)

    add_body(doc, "说明：天气、市场和日历的数据提供方将在实现阶段根据公司网络、地区可用性、许可条款和所需指数范围确定。", italic=True)

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())
