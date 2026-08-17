from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

fonts = [
    "Arial Unicode MS",
    "PingFang SC",
    "Hiragino Sans GB",
    "Heiti SC",
    "Songti SC",
    "STHeiti",
    "SimSong",
]

doc = Document()
for name in fonts:
    p = doc.add_paragraph()
    r = p.add_run(f"{name}：中文测试 桌面智能控制台 天气 股票 系统")
    r.font.name = name
    r.font.size = Pt(20)
    r._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    r._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    r._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    r._element.get_or_add_rPr().rFonts.set(qn("w:cs"), name)

doc.save("/Users/dongyu/ClaudeCode/ESP32/.docx_qa/font_test.docx")
